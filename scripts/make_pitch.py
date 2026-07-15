"""
Generate IdentiFind Sales Pitch & Investor Elevator Speech (.docx)
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Brand colors ────────────────────────────────────────────────────────────
NAVY   = RGBColor(0x0F, 0x17, 0x2A)   # #0F172A – dark background
BLUE   = RGBColor(0x1E, 0x40, 0xAF)   # #1E40AF – primary accent (deeper)
ACCENT = RGBColor(0x3B, 0x82, 0xF6)   # #3B82F6 – mid-blue
SLATE  = RGBColor(0x33, 0x4E, 0x68)   # slate heading
MUTED  = RGBColor(0x64, 0x74, 0x8B)   # muted / subtext
BLACK  = RGBColor(0x0F, 0x17, 0x2A)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_run(para, text, bold=False, italic=False, size=11, color=None, underline=False):
    run = para.add_run(text)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return run

def add_para(doc, text='', style='Normal', alignment=WD_ALIGN_PARAGRAPH.LEFT,
             space_before=0, space_after=6):
    p = doc.add_paragraph(style=style)
    p.alignment = alignment
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        p.add_run(text)
    return p

def add_section_heading(doc, text, space_before=16):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(4)
    # Bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '3B82F6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size = Pt(9)
    run.font.color.rgb = ACCENT
    run.font.name = 'Arial'
    return p

def add_bullet(doc, text, bold_prefix=None, indent_level=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_after = Pt(3)
    left_indent = Inches(0.25 + indent_level * 0.25)
    p.paragraph_format.left_indent = left_indent
    if bold_prefix:
        r = p.add_run(bold_prefix + ' ')
        r.bold = True
        r.font.size = Pt(10.5)
        r.font.name = 'Arial'
    r2 = p.add_run(text)
    r2.font.size = Pt(10.5)
    r2.font.name = 'Arial'
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CBD5E1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ════════════════════════════════════════════════════════════════════════════
doc = Document()

# ── Page setup: US Letter, 1" margins ───────────────────────────────────────
section = doc.sections[0]
section.page_width  = Inches(8.5)
section.page_height = Inches(11)
section.left_margin = section.right_margin = Inches(1)
section.top_margin = section.bottom_margin = Inches(1)

# ── Default style ────────────────────────────────────────────────────────────
style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)

# ═══════════════════════════════════════════════════════════════════════════
# COVER / TITLE BLOCK
# ═══════════════════════════════════════════════════════════════════════════
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.LEFT
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after = Pt(2)
t1 = title.add_run('IdentiFind')
t1.font.name = 'Arial'
t1.font.size = Pt(36)
t1.bold = True
t1.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

tagline = doc.add_paragraph()
tagline.alignment = WD_ALIGN_PARAGRAPH.LEFT
tagline.paragraph_format.space_before = Pt(0)
tagline.paragraph_format.space_after = Pt(2)
tl = tagline.add_run('Know What the Internet Knows About You.')
tl.font.name = 'Arial'
tl.font.size = Pt(16)
tl.italic = True
tl.font.color.rgb = ACCENT

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(12)
sr = sub.add_run('Investor Pitch & Elevator Speech  |  Confidential')
sr.font.size = Pt(9)
sr.font.color.rgb = MUTED
sr.font.name = 'Arial'

add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 1 — ELEVATOR SPEECH
# ═══════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '1.  30-Second Elevator Speech')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(6)
r = p.add_run(
    'Right now, your home address, phone number, family members’ names, '
    'and old passwords are sitting on data broker websites, breach databases, '
    'and social platforms — and you have no idea who’s looking at them.\n\n'
    'IdentiFind is the identity monitoring platform that finds you before the bad actors do. '
    'We scan hundreds of data brokers, breach databases, and 16+ social networks, '
    'then surface every exposure in a single risk score with clear, actionable steps to fix it. '
    'Think of it as a credit score for your digital privacy — one number that tells you '
    'exactly how exposed you are and what to do about it.\n\n'
    'We’re launching free for individuals and growing into a Pro subscription with '
    'real-time mobile alerts and continuous monitoring. '
    'The identity-protection market is worth $13 billion, growing 15% per year, '
    'and it’s full of clunky, expensive tools that don’t actually show you the problem. '
    'We do.'
)
r.font.size = Pt(11)
r.font.name = 'Arial'
r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

# Callout box — use a 1-cell table for the shaded quote
table = doc.add_table(rows=1, cols=1)
table.style = 'Table Grid'
cell = table.cell(0, 0)
set_cell_bg(cell, 'EFF6FF')
cell.width = Inches(6.5)
cp = cell.paragraphs[0]
cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
cp.paragraph_format.space_before = Pt(6)
cp.paragraph_format.space_after = Pt(6)
cr = cp.add_run(
    '“A credit score for your digital privacy.”'
)
cr.bold = True
cr.italic = True
cr.font.size = Pt(13)
cr.font.color.rgb = BLUE
cr.font.name = 'Arial'

doc.add_paragraph()  # spacer

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 2 — THE PROBLEM
# ═══════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '2.  The Problem')

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(8)
p2.paragraph_format.space_after = Pt(6)
r2 = p2.add_run(
    'Personal data is everywhere — and none of it is under your control. '
    'Every year, hundreds of data brokers legally buy and resell detailed profiles '
    'on nearly every American: full name, address history, phone numbers, relatives, '
    'income estimates, political affiliation, and more. '
    'At the same time, 3+ billion records are leaked in breaches annually, '
    'landing on dark-web forums where anyone can search them for free. '
    'Impersonators create fake social accounts, bad actors register lookalike domains, '
    'and the average victim doesn’t find out until real damage is done.'
)
r2.font.size = Pt(10.5)
r2.font.name = 'Arial'

add_bullet(doc, '33 million Americans are victims of identity fraud each year (Javelin, 2024).', bold_prefix='•')
add_bullet(doc, 'The average victim loses $1,100 and 200+ hours resolving it.', bold_prefix='•')
add_bullet(doc, 'Data brokers publish opt-out processes that are intentionally slow and fragmented — most people never complete them.', bold_prefix='•')
add_bullet(doc, 'Existing tools (LifeLock, Aura) focus on reactive fraud alerts, not proactive exposure discovery.', bold_prefix='•')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 3 — THE SOLUTION
# ═══════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '3.  The Solution — IdentiFind')

p3 = doc.add_paragraph()
p3.paragraph_format.space_before = Pt(8)
p3.paragraph_format.space_after = Pt(6)
r3 = p3.add_run(
    'IdentiFind gives individuals a single, unified view of their digital identity exposure '
    'across every major threat surface — all in real time.'
)
r3.font.size = Pt(10.5)
r3.font.name = 'Arial'

# Feature table
feature_data = [
    ('Breach Detection',        'Scans HIBP, DeHashed, IntelX, and Leak-Lookup for leaked credentials and pastes tied to your email or username.'),
    ('PII Broker Scanning',     'Queries Whitepages, Spokeo, BeenVerified, People Data Labs, FullContact, and others to surface what data brokers are publishing about you.'),
    ('Impersonation Scanner',   'Checks 16+ social platforms (GitHub, Twitter/X, Instagram, Reddit, TikTok, and more) for accounts using your name or handle variants.'),
    ('Lookalike Domain Watch',  'Monitors certificate transparency logs and DNS registrations for domains spoofing your name or brand.'),
    ('Social Account Audit',    'Reviews your connected accounts for missing MFA, unverified emails, and security gaps.'),
    ('Risk Score Engine',       'Every finding is classified CONFIRMED / HIGH / MEDIUM / LOW / SPECULATIVE with a 0–100 composite risk score and prioritized action steps.'),
]

tbl = doc.add_table(rows=1, cols=2)
tbl.style = 'Table Grid'
hdr_cells = tbl.rows[0].cells
set_cell_bg(hdr_cells[0], '1E40AF')
set_cell_bg(hdr_cells[1], '1E40AF')
for cell, txt in zip(hdr_cells, ['Feature', 'What It Does']):
    p = cell.paragraphs[0]
    r = p.add_run(txt)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = WHITE
    r.font.name = 'Arial'
    cell.width = Inches(1.9) if txt == 'Feature' else Inches(4.6)

for feat, desc in feature_data:
    row = tbl.add_row()
    set_cell_bg(row.cells[0], 'F8FAFC')
    pr = row.cells[0].paragraphs[0].add_run(feat)
    pr.bold = True
    pr.font.size = Pt(10)
    pr.font.name = 'Arial'
    pr.font.color.rgb = BLUE
    row.cells[0].width = Inches(1.9)
    dr = row.cells[1].paragraphs[0].add_run(desc)
    dr.font.size = Pt(10)
    dr.font.name = 'Arial'
    row.cells[1].width = Inches(4.6)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 4 — MARKET OPPORTUNITY
# ═══════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '4.  Market Opportunity')

p4 = doc.add_paragraph()
p4.paragraph_format.space_before = Pt(8)
p4.paragraph_format.space_after = Pt(6)
r4 = p4.add_run(
    'Identity protection is one of the fastest-growing segments in consumer security. '
    'Driven by a wave of high-profile breaches, state privacy laws (CCPA, CPRA), '
    'and rising awareness of data broker abuse, the market is expected to exceed '
    '$25 billion by 2030.'
)
r4.font.size = Pt(10.5)
r4.font.name = 'Arial'

mkt_data = [
    ('Total Addressable Market (TAM)', '$13B+ (2024)',    'Global identity protection & monitoring market.'),
    ('Serviceable Addressable Market', '$4.2B',          'US adults with digital privacy concerns & willingness to pay.'),
    ('Serviceable Obtainable Market',  '$42M (Year 3)',   'Conservative 1% share of SAM with 85,000 paid subscribers.'),
    ('Growth Rate',                    '~15% CAGR',      'Driven by breach volume, regulation, and public awareness.'),
]

mt = doc.add_table(rows=1, cols=3)
mt.style = 'Table Grid'
for cell, txt in zip(mt.rows[0].cells, ['Segment', 'Size', 'Notes']):
    set_cell_bg(cell, '0F172A')
    r = cell.paragraphs[0].add_run(txt)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = WHITE
    r.font.name = 'Arial'
mt.rows[0].cells[0].width = Inches(2.5)
mt.rows[0].cells[1].width = Inches(1.3)
mt.rows[0].cells[2].width = Inches(2.7)

for label, size, note in mkt_data:
    row = mt.add_row()
    row.cells[0].width = Inches(2.5)
    row.cells[1].width = Inches(1.3)
    row.cells[2].width = Inches(2.7)
    row.cells[0].paragraphs[0].add_run(label).font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].font.name = 'Arial'
    sz = row.cells[1].paragraphs[0].add_run(size)
    sz.bold = True
    sz.font.size = Pt(10)
    sz.font.color.rgb = BLUE
    sz.font.name = 'Arial'
    row.cells[2].paragraphs[0].add_run(note).font.size = Pt(10)
    row.cells[2].paragraphs[0].runs[0].font.name = 'Arial'

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 5 — BUSINESS MODEL
# ═══════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '5.  Business Model')

p5 = doc.add_paragraph()
p5.paragraph_format.space_before = Pt(8)
p5.paragraph_format.space_after = Pt(6)
r5 = p5.add_run(
    'IdentiFind follows a proven freemium-to-subscription model designed to maximize '
    'top-of-funnel growth while converting high-value users to recurring revenue.'
)
r5.font.size = Pt(10.5)
r5.font.name = 'Arial'

tier_data = [
    ('Free',          '$0 / mo',  'Monthly scan, risk score, top 5 findings, breach check.',           'Acquisition & awareness'),
    ('Pro',           '$12 / mo', 'Continuous monitoring, all findings, real-time mobile push alerts, data broker reports, priority support.', 'Core revenue driver'),
    ('Family',        '$22 / mo', 'Up to 5 profiles, all Pro features, shared dashboard.',             'Expansion revenue'),
]

tt = doc.add_table(rows=1, cols=4)
tt.style = 'Table Grid'
for cell, txt in zip(tt.rows[0].cells, ['Tier', 'Price', 'Includes', 'Goal']):
    set_cell_bg(cell, '1E40AF')
    r = cell.paragraphs[0].add_run(txt)
    r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = WHITE
    r.font.name = 'Arial'

widths = [Inches(0.9), Inches(0.9), Inches(3.5), Inches(1.7)]
for i, w in enumerate(widths):
    tt.rows[0].cells[i].width = w

for tier, price, inc, goal in tier_data:
    row = tt.add_row()
    for i, w in enumerate(widths):
        row.cells[i].width = w
    pr = row.cells[0].paragraphs[0].add_run(tier)
    pr.bold = True; pr.font.size = Pt(10); pr.font.name = 'Arial'; pr.font.color.rgb = BLUE
    row.cells[1].paragraphs[0].add_run(price).font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].font.name = 'Arial'
    row.cells[1].paragraphs[0].runs[0].bold = True
    row.cells[2].paragraphs[0].add_run(inc).font.size = Pt(10)
    row.cells[2].paragraphs[0].runs[0].font.name = 'Arial'
    row.cells[3].paragraphs[0].add_run(goal).font.size = Pt(10)
    row.cells[3].paragraphs[0].runs[0].font.name = 'Arial'
    row.cells[3].paragraphs[0].runs[0].italic = True

doc.add_paragraph()

p5b = doc.add_paragraph()
p5b.paragraph_format.space_after = Pt(4)
r5b = p5b.add_run('Revenue projection (conservative):')
r5b.bold = True
r5b.font.size = Pt(10.5)
r5b.font.name = 'Arial'

add_bullet(doc, 'Year 1: 5,000 paid subscribers → ~$720K ARR', bold_prefix='•')
add_bullet(doc, 'Year 2: 25,000 paid subscribers → ~$3.6M ARR', bold_prefix='•')
add_bullet(doc, 'Year 3: 85,000 paid subscribers → ~$12.2M ARR', bold_prefix='•')
add_bullet(doc, 'Target LTV:CAC ratio of 4:1 via SEO, content, and referral loops.', bold_prefix='•')

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 6 — COMPETITIVE DIFFERENTIATION
# ═══════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '6.  Why IdentiFind Wins')

p6 = doc.add_paragraph()
p6.paragraph_format.space_before = Pt(8)
p6.paragraph_format.space_after = Pt(6)
r6 = p6.add_run(
    'Incumbents like LifeLock and Aura are expensive, opaque, and reactive — '
    'they tell you after the damage is done. DeleteMe and Kanary remove data '
    'but don’t show you the full picture. IdentiFind is the only platform that '
    'combines discovery, scoring, and remediation guidance in a single, affordable product.'
)
r6.font.size = Pt(10.5)
r6.font.name = 'Arial'

comp_data = [
    ('Feature',                    'IdentiFind', 'LifeLock', 'Aura',    'DeleteMe'),
    ('Real-time risk score',       'Yes',         'No',       'Partial', 'No'),
    ('Data broker scanning',       'Yes',         'Partial',  'Yes',     'Yes'),
    ('Impersonation detection',    'Yes',         'No',       'No',      'No'),
    ('Lookalike domain watch',     'Yes',         'No',       'No',      'No'),
    ('Social account security',    'Yes',         'No',       'No',      'No'),
    ('Mobile app with push alerts','Yes (Q4 \'26)','Yes',     'Yes',     'No'),
    ('Price (entry)',               'Free',        '$11.99/mo','$3/mo',  '$9.99/mo'),
]

ct = doc.add_table(rows=len(comp_data), cols=5)
ct.style = 'Table Grid'
for ci, row_data in enumerate(comp_data):
    for cj, val in enumerate(row_data):
        cell = ct.rows[ci].cells[cj]
        if ci == 0:
            set_cell_bg(cell, '0F172A')
            r = cell.paragraphs[0].add_run(val)
            r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(9.5); r.font.name = 'Arial'
        else:
            if cj == 1:  # IdentiFind column
                set_cell_bg(cell, 'EFF6FF')
            r = cell.paragraphs[0].add_run(val)
            r.font.size = Pt(9.5); r.font.name = 'Arial'
            if cj == 1:
                r.bold = True; r.font.color.rgb = BLUE
            if val == 'Yes':
                r.font.color.rgb = RGBColor(0x16, 0xA3, 0x4A)
            elif val == 'No':
                r.font.color.rgb = RGBColor(0xDC, 0x26, 0x26)

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 7 — TRACTION & ROADMAP
# ═══════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '7.  Where We Are')

p7 = doc.add_paragraph()
p7.paragraph_format.space_before = Pt(8)
p7.paragraph_format.space_after = Pt(6)
r7 = p7.add_run('Current status (June 2026):')
r7.bold = True
r7.font.size = Pt(10.5)
r7.font.name = 'Arial'

add_bullet(doc, 'Full-stack web platform live: Next.js + Supabase + Prisma, deployed and functional.', bold_prefix='✓')
add_bullet(doc, '7-phase scan pipeline complete: breach detection, PII broker queries, impersonation scanner, lookalike domain watch, confidence scoring, and DB persistence.', bold_prefix='✓')
add_bullet(doc, 'OAuth login with 5 providers (Google, GitHub, LinkedIn, Twitter/X, Facebook).', bold_prefix='✓')
add_bullet(doc, 'AES-256-GCM encryption for all stored PII.', bold_prefix='✓')
add_bullet(doc, 'React Native / Expo mobile scaffold complete — Phase 1 done, targeting app store submission Q4 2026.', bold_prefix='✓')
add_bullet(doc, 'Live API integrations: IntelX, Hunter.io, crt.sh; HIBP and People Data Labs queued.', bold_prefix='✓')

p7b = doc.add_paragraph()
p7b.paragraph_format.space_before = Pt(8)
p7b.paragraph_format.space_after = Pt(4)
r7b = p7b.add_run('Near-term milestones:')
r7b.bold = True
r7b.font.size = Pt(10.5)
r7b.font.name = 'Arial'

roadmap = [
    ('Q3 2026', 'Public beta launch, HIBP + PDL API integrations, user onboarding flow.'),
    ('Q4 2026', 'iOS & Android app store launch, push notifications, Pro tier activation.'),
    ('Q1 2027', 'Family plan, removal request automation, affiliate & SEO growth push.'),
    ('Q2 2027', 'Series A / strategic partnerships, enterprise pilot (HR teams, law firms).'),
]

rt = doc.add_table(rows=1, cols=2)
rt.style = 'Table Grid'
for cell, txt in zip(rt.rows[0].cells, ['Quarter', 'Milestone']):
    set_cell_bg(cell, '334E68')
    r = cell.paragraphs[0].add_run(txt)
    r.bold = True; r.font.color.rgb = WHITE; r.font.size = Pt(10); r.font.name = 'Arial'
rt.rows[0].cells[0].width = Inches(1.2)
rt.rows[0].cells[1].width = Inches(5.3)

for q, mile in roadmap:
    row = rt.add_row()
    row.cells[0].width = Inches(1.2)
    row.cells[1].width = Inches(5.3)
    qr = row.cells[0].paragraphs[0].add_run(q)
    qr.bold = True; qr.font.size = Pt(10); qr.font.name = 'Arial'; qr.font.color.rgb = BLUE
    mr = row.cells[1].paragraphs[0].add_run(mile)
    mr.font.size = Pt(10); mr.font.name = 'Arial'

doc.add_paragraph()

# ═══════════════════════════════════════════════════════════════════════════
# SECTION 8 — THE ASK
# ═══════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '8.  The Ask')

p8 = doc.add_paragraph()
p8.paragraph_format.space_before = Pt(8)
p8.paragraph_format.space_after = Pt(6)
r8 = p8.add_run(
    'We are raising a pre-seed round to fund the following over the next 18 months:'
)
r8.font.size = Pt(10.5)
r8.font.name = 'Arial'

add_bullet(doc, 'API integration budget: HIBP, People Data Labs, DeHashed, and Pipl subscriptions to fully activate the scan pipeline.', bold_prefix='•')
add_bullet(doc, 'Mobile app completion: Phases 2–5 (core screens, push notifications, biometrics, App Store submission).', bold_prefix='•')
add_bullet(doc, 'Marketing & growth: SEO content, paid acquisition tests, and referral program build-out.', bold_prefix='•')
add_bullet(doc, 'Infrastructure: Supabase scaling, CDN, security audit, and SOC 2 Type I preparation.', bold_prefix='•')

p8b = doc.add_paragraph()
p8b.paragraph_format.space_before = Pt(8)
p8b.paragraph_format.space_after = Pt(4)
r8b = p8b.add_run(
    'In return, we offer equity in a product with a working prototype, a complete technical foundation, '
    'and a clear path to $3M+ ARR within 24 months of funding.'
)
r8b.font.size = Pt(10.5)
r8b.font.name = 'Arial'
r8b.bold = True

add_divider(doc)

# ═══════════════════════════════════════════════════════════════════════════
# CLOSING / CONTACT
# ═══════════════════════════════════════════════════════════════════════════
p9 = doc.add_paragraph()
p9.alignment = WD_ALIGN_PARAGRAPH.CENTER
p9.paragraph_format.space_before = Pt(14)
p9.paragraph_format.space_after = Pt(4)
r9 = p9.add_run('Alexander Grant Freshley  |  Founder, IdentiFind')
r9.bold = True; r9.font.size = Pt(11); r9.font.name = 'Arial'

p10 = doc.add_paragraph()
p10.alignment = WD_ALIGN_PARAGRAPH.CENTER
p10.paragraph_format.space_after = Pt(2)
r10 = p10.add_run('alexander.freshley@gmail.com')
r10.font.size = Pt(10.5); r10.font.name = 'Arial'; r10.font.color.rgb = ACCENT

out = '/sessions/keen-nifty-tesla/mnt/IdentiFind/IdentiFind Investor Pitch.docx'
doc.save(out)
print(f'Saved to {out}')
