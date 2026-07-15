"""
IdentiFind Mobile App — Planning Document Generator
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ─── Page setup ───────────────────────────────────────────────────────────────
section = doc.sections[0]
section.page_width    = Inches(8.5)
section.page_height   = Inches(11)
section.top_margin    = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin   = Inches(1)
section.right_margin  = Inches(1)

# ─── Colors ───────────────────────────────────────────────────────────────────
DARK   = RGBColor(0x0F, 0x17, 0x2A)
ACCENT = RGBColor(0x3B, 0x82, 0xF6)
MUTED  = RGBColor(0x64, 0x74, 0x8B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)
GREEN  = RGBColor(0x16, 0xA3, 0x4A)

# ─── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="CCCCCC"):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for side in ['top', 'left', 'bottom', 'right']:
        border = OxmlElement(f'w:{side}')
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '4')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), color)
        tcBorders.append(border)
    tcPr.append(tcBorders)

def spacer(pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    p.add_run('').font.size = Pt(pt)

def body(text, size=10.5, bold=False, italic=False, color=None, sb=0, sa=6, align=WD_ALIGN_PARAGRAPH.LEFT, indent=0):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(sb)
    p.paragraph_format.space_after  = Pt(sa)
    p.paragraph_format.left_indent  = Inches(indent)
    run = p.add_run(text)
    run.font.size   = Pt(size)
    run.font.bold   = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    return p

def h1(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text.upper())
    run.font.size  = Pt(10)
    run.font.bold  = True
    run.font.color.rgb = ACCENT
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '4')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '3B82F6')
    pBdr.append(bottom)
    pPr.append(pBdr)

def h2(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.size  = Pt(11.5)
    run.font.bold  = True
    run.font.color.rgb = DARK

def bul(text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Inches(0.25 + indent * 0.2)
    run = p.add_run(text)
    run.font.size = Pt(10)
    run.font.color.rgb = DARK

def table_header_row(table, headers, bg="0F172A"):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        set_cell_bg(cell, bg)
        set_cell_borders(cell, "1E293B")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run(h)
        run.font.size  = Pt(9)
        run.font.bold  = True
        run.font.color.rgb = WHITE

def table_data_row(table, data, bg="F8FAFC", alt_bg="FFFFFF", ri=0):
    row = table.add_row()
    bg_use = bg if ri % 2 == 0 else alt_bg
    for ci, text in enumerate(data):
        cell = row.cells[ci]
        set_cell_bg(cell, bg_use)
        set_cell_borders(cell, "E2E8F0")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(3)
        run = p.add_run(str(text))
        run.font.size = Pt(9.5)
        run.font.color.rgb = GREEN if "NEW" in str(text) else (MUTED if str(text) == "—" else DARK)

# ─── Cover ────────────────────────────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.LEFT
p.paragraph_format.space_before = Pt(36)
p.paragraph_format.space_after  = Pt(0)
r = p.add_run("IdentiFind")
r.font.size = Pt(30); r.font.bold = True; r.font.color.rgb = DARK

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(2)
p2.paragraph_format.space_after  = Pt(0)
r2 = p2.add_run("Mobile Application — Development Plan")
r2.font.size = Pt(16); r2.font.color.rgb = ACCENT

spacer(10)

# Meta table
mt = doc.add_table(rows=1, cols=4)
mt.style = 'Table Grid'
for i, (label, val) in enumerate([("Version","1.0"),("Date","May 2026"),("Target Release","Q4 2026"),("Platforms","iOS & Android")]):
    cell = mt.rows[0].cells[i]
    set_cell_bg(cell, "F1F5F9")
    set_cell_borders(cell, "E2E8F0")
    p = cell.paragraphs[0]
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(4)
    lr = p.add_run(label + "\n"); lr.font.size = Pt(8); lr.font.bold = True; lr.font.color.rgb = MUTED
    vr = p.add_run(val);         vr.font.size = Pt(10); vr.font.bold = True; vr.font.color.rgb = DARK

spacer(20)

# ─── 1. Overview ──────────────────────────────────────────────────────────────
h1("1. Overview")
spacer(4)
body(
    "This document outlines the plan to develop native iOS and Android mobile applications for IdentiFind, "
    "the identity monitoring and PII discovery platform. The mobile apps will surface all existing web "
    "features — risk scores, security findings, breach alerts, account management, and PII exposure reports — "
    "optimized for mobile, with the addition of real-time push notifications for critical security events.",
    sa=6
)
body(
    "The existing backend (Next.js API routes, Prisma/Supabase database, scan pipeline) requires minimal "
    "changes. The mobile application is a new client consuming the existing API, meaning the scan engine, "
    "breach detectors, PII brokers, and security scoring logic ship to mobile automatically.",
    sa=6
)

# ─── 2. Stack ─────────────────────────────────────────────────────────────────
h1("2. Technology Stack")
spacer(4)
h2("Framework: React Native with Expo")
body("React Native with Expo is the recommended framework for the following reasons:", sa=4)
bul("Existing codebase is React/TypeScript — component patterns, hooks, and types transfer directly")
bul("TypeScript types (PiiRecord, BreachRecord, ScoredFinding, etc.) are shared between web and mobile with no duplication")
bul("Expo handles build pipelines, app store submission (EAS Build), and push notification delivery (Expo Notifications)")
bul("Expo Router uses the same file-system routing paradigm as Next.js App Router — consistent mental model")
bul("Single codebase targets both iOS and Android; no Swift/Kotlin split")
spacer(6)

h2("Alternatives Considered")
at = doc.add_table(rows=1, cols=3)
at.style = 'Table Grid'
table_header_row(at, ["Option", "Pro", "Con"])
alts = [
    ["Flutter / Dart",        "High performance, good UI",           "No code sharing with existing TS codebase; team must learn Dart"],
    ["Native Swift + Kotlin", "Best performance, full platform APIs", "Two separate codebases; 2x maintenance; longest timeline"],
    ["PWA (web wrapper)",     "Zero new codebase, instant deploy",    "iOS push notifications unreliable; no biometric auth; app store rejection risk"],
    ["Capacitor / Ionic",     "Wraps existing Next.js app quickly",   "Web performance in native shell; complex OAuth redirects"],
]
for ri, row in enumerate(alts):
    table_data_row(at, row, ri=ri)
spacer(8)

# ─── 3. Architecture ──────────────────────────────────────────────────────────
h1("3. Architecture")
spacer(4)
h2("What Stays the Same")
body("The entire backend is reused without modification:", sa=4)
bul("Next.js API routes: /api/scan, /api/alerts, /api/accounts, /api/settings/*")
bul("Scan orchestrator and all scanner modules (breach, PII broker, username, crt.sh)")
bul("Prisma schema and Supabase PostgreSQL database")
bul("NextAuth session management — mobile uses the same auth endpoints via browser handoff")
bul("Encryption, audit logging, and security engine")
spacer(6)

h2("New Backend Additions (Minimal)")
bul("Push token endpoint — POST /api/notifications/register stores an Expo push token per user session")
bul("Notification dispatch — added to scan orchestrator: fires a push when CRITICAL or HIGH findings are returned")
bul("Scheduled scan trigger — Vercel Cron Job that auto-triggers daily scans per user for background alerts")
spacer(6)

h2("Mobile Client Architecture")
bul("Expo Router for file-system navigation (tabs: Dashboard, Alerts, Accounts, Profile)")
bul("React Query for data fetching, caching, and background refetch from the Next.js API")
bul("expo-secure-store for encrypted session token storage (replaces browser cookies)")
bul("expo-notifications for push notification registration and foreground/background handling")
bul("expo-local-authentication for biometric lock screen (Face ID / fingerprint)")
bul("expo-auth-session for OAuth provider flows (Google, GitHub, LinkedIn, Twitter, Facebook)")
spacer(8)

# ─── 4. Feature Parity ────────────────────────────────────────────────────────
h1("4. Feature Parity")
spacer(4)
body("All existing web features are included. Mobile adds three new capabilities:", sa=6)

ft = doc.add_table(rows=1, cols=3)
ft.style = 'Table Grid'
table_header_row(ft, ["Feature", "Web", "Mobile"])
feats = [
    ["Risk score dashboard",          "✓", "✓"],
    ["Security findings list",        "✓", "✓"],
    ["Connected accounts management", "✓", "✓"],
    ["Breach alerts",                 "✓", "✓"],
    ["Impersonation alerts",          "✓", "✓"],
    ["PII exposure report",           "✓", "✓"],
    ["Manual scan trigger",           "✓", "✓"],
    ["Profile & settings",            "✓", "✓"],
    ["Push notifications",            "—", "✓  (NEW)"],
    ["Biometric lock (Face ID / FP)", "—", "✓  (NEW)"],
    ["Background / scheduled scans",  "—", "✓  (NEW)"],
]
for ri, f in enumerate(feats):
    table_data_row(ft, f, ri=ri)
spacer(8)

# ─── 5. Phases ────────────────────────────────────────────────────────────────
h1("5. Development Phases")
spacer(4)

phases = [
    ("Phase 1", "Foundation",             "June 2026",       "1E40AF",
     ["Initialize Expo project with TypeScript, Expo Router, and EAS Build configuration",
      "Implement OAuth login flow via expo-auth-session connecting to existing NextAuth backend",
      "Build tab navigation shell: Dashboard, Alerts, Accounts, Profile",
      "Wire up React Query to the existing Next.js API",
      "Implement expo-secure-store for session token persistence",
      "Verify full auth round-trip on iOS Simulator and Android Emulator",
      "Goal: authenticated user can log in and see their live data on both platforms"]),
    ("Phase 2", "Core Screens",           "July – Aug 2026", "1D4ED8",
     ["Dashboard: risk score ring, finding counts, last-scanned timestamp, pull-to-refresh scan trigger",
      "Alerts screen: breach, impersonation, and PII findings with severity badges",
      "Alert detail: full description, recommended action, resolve button",
      "Accounts screen: connected social accounts list, add/remove account flow",
      "Profile screen: name, email, phone verification status",
      "Settings: notification preferences, scan frequency, biometric toggle, sign out",
      "Loading skeletons and empty states for all screens"]),
    ("Phase 3", "Push Notifications",     "September 2026",  "2563EB",
     ["Add POST /api/notifications/register to store Expo push tokens",
      "Register for push permissions on app launch; send token to backend",
      "Add notification dispatch to scan orchestrator for CRITICAL/HIGH findings",
      "Set up Vercel Cron Job for daily automated scans per active user",
      "Handle foreground notification display and navigation to relevant alert",
      "Handle background notification tap with deep-link to specific finding",
      "Test full push delivery loop on physical devices (required for push testing)"]),
    ("Phase 4", "Polish & Security",      "October 2026",    "3B82F6",
     ["Biometric authentication gate on app open via expo-local-authentication",
      "Auto-lock after configurable inactivity period",
      "Haptic feedback on critical alerts",
      "Dark mode support matching the existing brand palette",
      "Accessibility audit (VoiceOver / TalkBack) on all screens",
      "Error boundary and crash reporting (Sentry via @sentry/react-native)",
      "Offline state handling — graceful degradation when network is unavailable"]),
    ("Phase 5", "Beta & Submission",      "November 2026",   "60A5FA",
     ["Internal TestFlight build for iOS; Internal Testing track for Android",
      "Structured QA pass across iPhone and Android physical devices",
      "App Store Connect: screenshots, description, privacy policy, data usage labels",
      "Google Play Console: store graphics, content rating, data safety section",
      "Submit for review (Apple: 1–7 days typical; Google: 1–3 days typical)",
      "Address any rejection feedback and resubmit"]),
    ("Phase 6", "Launch Buffer",          "December 2026",   "93C5FD",
     ["Monitor crash reports and ANRs from production users",
      "Patch high-priority issues in rapid follow-up builds",
      "Collect early user feedback on notification frequency and UX",
      "Plan v1.1 backlog: Reddit/Discord/TikTok clients, iOS home screen widget, watchOS complication"]),
]

for num, title, month, color, items in phases:
    pht = doc.add_table(rows=1, cols=2)
    pht.style = 'Table Grid'
    lc = pht.rows[0].cells[0]
    rc = pht.rows[0].cells[1]
    set_cell_bg(lc, color); set_cell_borders(lc, color)
    set_cell_bg(rc, color); set_cell_borders(rc, color)
    lp = lc.paragraphs[0]; lp.paragraph_format.space_before = Pt(5); lp.paragraph_format.space_after = Pt(5)
    lr = lp.add_run(num); lr.font.size = Pt(9.5); lr.font.bold = True; lr.font.color.rgb = WHITE
    rp = rc.paragraphs[0]; rp.paragraph_format.space_before = Pt(5); rp.paragraph_format.space_after = Pt(5)
    rr1 = rp.add_run(title); rr1.font.size = Pt(9.5); rr1.font.bold = True; rr1.font.color.rgb = WHITE
    rr2 = rp.add_run(f"   {month}"); rr2.font.size = Pt(9); rr2.font.color.rgb = RGBColor(0xBF, 0xDB, 0xFE)
    for item in items:
        bul(item)
    spacer(6)

# ─── 6. Timeline ──────────────────────────────────────────────────────────────
h1("6. Timeline Summary")
spacer(4)
tlt = doc.add_table(rows=1, cols=4)
tlt.style = 'Table Grid'
table_header_row(tlt, ["Phase", "Focus", "Month", "Milestone"])
tl_rows = [
    ["1", "Foundation",        "June",         "Auth works on both simulators"],
    ["2", "Core Screens",      "July – Aug",   "All web features visible on mobile"],
    ["3", "Push Notifications","September",    "Alerts firing on physical devices"],
    ["4", "Polish",            "October",      "Biometrics, dark mode, accessibility"],
    ["5", "Beta & Submission", "November",     "Live on TestFlight + Play Internal Testing"],
    ["6", "Launch Buffer",     "December",     "v1.0 public — monitor & patch"],
]
for ri, row in enumerate(tl_rows):
    table_data_row(tlt, row, ri=ri)
spacer(8)

# ─── 7. Risks ─────────────────────────────────────────────────────────────────
h1("7. Risks & Mitigations")
spacer(4)
risks = [
    ("OAuth flow on mobile",
     "Mobile OAuth requires browser redirects and deep-link callbacks — more complex than web. "
     "Mitigation: expo-auth-session handles this pattern; implement and validate in Phase 1 before any other work."),
    ("Apple App Store review",
     "Apps handling personal identity data may require additional privacy review. "
     "Mitigation: prepare a clear privacy policy, accurate data-usage labels, and a demo account for reviewers. Allow 2–3 week buffer in Phase 5."),
    ("Push notification opt-in rates",
     "iOS requires explicit user permission; some users will decline. "
     "Mitigation: request permission contextually after a scan finds something, not on first launch."),
    ("Expo managed vs. bare workflow",
     "Some native modules may require ejecting from the Expo managed workflow. "
     "Mitigation: audit all required native modules before scaffolding; start with bare workflow if any module requires it."),
    ("iOS background execution limits",
     "iOS heavily throttles background app refresh — full scans cannot reliably run on the client. "
     "Mitigation: run scans server-side via cron and push results to the device. No background execution required on the client."),
]
for title, desc in risks:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5); p.paragraph_format.space_after = Pt(1)
    r1 = p.add_run(title); r1.font.size = Pt(10.5); r1.font.bold = True; r1.font.color.rgb = DARK
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(5)
    p2.paragraph_format.left_indent = Inches(0.2)
    r2 = p2.add_run(desc); r2.font.size = Pt(10); r2.font.color.rgb = MUTED
spacer(8)

# ─── 8. Phase 1 Technical Detail ──────────────────────────────────────────────
h1("8. Phase 1 Scaffold — Technical Detail")
spacer(4)
body(
    "Completion is defined as: an authenticated user can log in via OAuth on both an iOS Simulator "
    "and Android Emulator, navigate the tab shell, and see their live risk score and findings pulled from the existing API.",
    sa=8
)
p1 = [
    ("1. Initialize Expo project",
     "npx create-expo-app@latest identifind-mobile --template tabs\n"
     "Configure TypeScript strict mode, Expo Router, and path aliases matching the web project."),
    ("2. Install core dependencies",
     "expo-auth-session, expo-secure-store, expo-local-authentication, expo-notifications, "
     "@tanstack/react-query."),
    ("3. Configure EAS Build",
     "eas build:configure — sets up iOS and Android build profiles for simulator development builds."),
    ("4. Implement authentication",
     "expo-auth-session initiates OAuth flows; NextAuth /api/auth/* endpoints handle tokens unchanged. "
     "Session token is persisted to expo-secure-store."),
    ("5. Build navigation shell",
     "Four tabs: Dashboard, Alerts, Accounts, Profile. Each screen fetches from the existing Next.js API "
     "using the stored session token."),
    ("6. Dashboard data fetch",
     "GET /api/scan retrieves the last scan result. Render risk score and finding counts. "
     "Pull-to-refresh triggers POST /api/scan and refetches."),
    ("7. Validate on both platforms",
     "Run on iOS Simulator and Android Emulator. Confirm auth, navigation, and data fetch work end-to-end."),
]
for step_title, step_desc in p1:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(1)
    r = p.add_run(step_title); r.font.size = Pt(10.5); r.font.bold = True; r.font.color.rgb = DARK
    p2 = doc.add_paragraph()
    p2.paragraph_format.space_before = Pt(0); p2.paragraph_format.space_after = Pt(4)
    p2.paragraph_format.left_indent = Inches(0.2)
    r2 = p2.add_run(step_desc); r2.font.size = Pt(10); r2.font.color.rgb = MUTED

spacer(16)
body("— End of Document —", size=9, color=MUTED, align=WD_ALIGN_PARAGRAPH.CENTER)

# ─── Save ─────────────────────────────────────────────────────────────────────
out = "/sessions/keen-nifty-tesla/mnt/IdentiFind/IdentiFind Mobile App Plan.docx"
doc.save(out)
print(f"Saved: {out}")
