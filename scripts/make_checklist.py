"""
Generate IdentiFind Launch Checklist (.docx)
Investor-readiness & professional operations checklist
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# Colors
NAVY   = RGBColor(0x0F, 0x17, 0x2A)
BLUE   = RGBColor(0x1E, 0x40, 0xAF)
ACCENT = RGBColor(0x3B, 0x82, 0xF6)
GREEN  = RGBColor(0x16, 0xA3, 0x4A)
AMBER  = RGBColor(0xD9, 0x77, 0x06)
RED    = RGBColor(0xDC, 0x26, 0x26)
MUTED  = RGBColor(0x64, 0x74, 0x8B)
WHITE  = RGBColor(0xFF, 0xFF, 0xFF)

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def add_section_heading(doc, text, space_before=14):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '6')
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), '3B82F6')
    pBdr.append(bottom)
    pPr.append(pBdr)
    run = p.add_run(text.upper())
    run.bold = True
    run.font.size  = Pt(9)
    run.font.color.rgb = ACCENT
    run.font.name  = 'Arial'
    return p

def add_divider(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'),   'single')
    bottom.set(qn('w:sz'),    '4')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'CBD5E1')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

PRIORITY_COLORS = {
    'CRITICAL': ('DC2626', WHITE),   # red bg, white text
    'HIGH':     ('D97706', WHITE),   # amber bg, white text
    'MEDIUM':   ('1E40AF', WHITE),   # blue bg, white text
    'LOW':      ('64748B', WHITE),   # slate bg, white text
    'DONE':     ('16A34A', WHITE),   # green bg, white text
}

def add_item_row(tbl, done, priority, item, detail, cost, link=''):
    row = tbl.add_row()
    # Col 0: status checkbox
    check = '✓' if done else '☐'
    cp = row.cells[0].paragraphs[0]
    cr = cp.add_run(check)
    cr.font.size = Pt(11)
    cr.font.name = 'Arial'
    cr.font.color.rgb = GREEN if done else MUTED
    cr.bold = done

    # Col 1: priority badge
    pri_hex, pri_txt_color = PRIORITY_COLORS.get(priority, ('64748B', WHITE))
    if done:
        pri_hex, pri_txt_color = PRIORITY_COLORS['DONE']
        priority = 'DONE'
    set_cell_bg(row.cells[1], pri_hex)
    pp = row.cells[1].paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pr = pp.add_run(priority)
    pr.bold = True
    pr.font.size = Pt(8)
    pr.font.name = 'Arial'
    pr.font.color.rgb = pri_txt_color

    # Col 2: item + detail
    ip = row.cells[2].paragraphs[0]
    ir = ip.add_run(item)
    ir.bold = True
    ir.font.size = Pt(10)
    ir.font.name = 'Arial'
    if detail:
        ip2 = row.cells[2].add_paragraph()
        ip2.paragraph_format.space_before = Pt(1)
        id2 = ip2.add_run(detail)
        id2.font.size = Pt(9)
        id2.font.name = 'Arial'
        id2.font.color.rgb = MUTED
        if link:
            id2.text = detail + '  →  ' + link

    # Col 3: cost
    cp2 = row.cells[3].paragraphs[0]
    cp2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cor = cp2.add_run(cost)
    cor.font.size = Pt(9.5)
    cor.font.name = 'Arial'
    if cost.startswith('Free') or cost == '—':
        cor.font.color.rgb = GREEN
    else:
        cor.font.color.rgb = BLUE

    # Row widths
    row.cells[0].width = Inches(0.35)
    row.cells[1].width = Inches(0.85)
    row.cells[2].width = Inches(5.1)
    row.cells[3].width = Inches(1.2)

    return row

def make_table(doc):
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = 'Table Grid'
    hdr = tbl.rows[0]
    set_cell_bg(hdr.cells[0], '0F172A')
    set_cell_bg(hdr.cells[1], '0F172A')
    set_cell_bg(hdr.cells[2], '0F172A')
    set_cell_bg(hdr.cells[3], '0F172A')
    for cell, txt in zip(hdr.cells, ['', 'Priority', 'Item', 'Cost']):
        r = cell.paragraphs[0].add_run(txt)
        r.bold = True; r.font.size = Pt(9.5); r.font.name = 'Arial'; r.font.color.rgb = WHITE
    hdr.cells[0].width = Inches(0.35)
    hdr.cells[1].width = Inches(0.85)
    hdr.cells[2].width = Inches(5.1)
    hdr.cells[3].width = Inches(1.2)
    return tbl


# ════════════════════════════════════════════════════════════════════════════
doc = Document()

section = doc.sections[0]
section.page_width   = Inches(8.5)
section.page_height  = Inches(11)
section.left_margin  = section.right_margin  = Inches(1)
section.top_margin   = section.bottom_margin = Inches(1)

style = doc.styles['Normal']
style.font.name = 'Arial'
style.font.size = Pt(10.5)

# ── Title ────────────────────────────────────────────────────────────────────
title = doc.add_paragraph()
title.paragraph_format.space_before = Pt(0)
title.paragraph_format.space_after  = Pt(2)
tr = title.add_run('IdentiFind — Launch Readiness Checklist')
tr.bold = True; tr.font.size = Pt(26); tr.font.name = 'Arial'
tr.font.color.rgb = NAVY

sub = doc.add_paragraph()
sub.paragraph_format.space_after = Pt(4)
sr = sub.add_run('Investor & Operational Readiness  |  Updated June 2026')
sr.font.size = Pt(9); sr.font.name = 'Arial'; sr.font.color.rgb = MUTED

# Legend
leg = doc.add_paragraph()
leg.paragraph_format.space_after = Pt(8)
for label, color in [('CRITICAL', RED), ('HIGH', AMBER), ('MEDIUM', BLUE), ('LOW', MUTED), ('DONE', GREEN)]:
    r = leg.add_run(f'  {label}  ')
    r.bold = True; r.font.size = Pt(8); r.font.name = 'Arial'; r.font.color.rgb = color
leg.add_run('   priority levels').font.size = Pt(9)

add_divider(doc)

# ════════════════════════════════════════════════════════════════════════════
# SECTION 1 — DOMAIN & IDENTITY
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '1.  Domain & Brand Identity')

p = doc.add_paragraph()
p.paragraph_format.space_before = Pt(6)
p.paragraph_format.space_after  = Pt(8)
pr = p.add_run(
    'The first thing an investor does after a meeting is search for you. '
    'These items must be in place before any external pitch or demo.'
)
pr.font.size = Pt(10); pr.font.name = 'Arial'; pr.font.color.rgb = MUTED

tbl1 = make_table(doc)
add_item_row(tbl1, False, 'CRITICAL', 'Register identifind.com (or best available TLD)',
    'Check Namecheap or Google Domains. If .com is taken, try .app or .io — both are '
    'credible for software products. Avoid .net or hyphens.',
    '~$12–20/yr', 'namecheap.com / domains.google')
add_item_row(tbl1, False, 'CRITICAL', 'Set up Google Workspace for professional email',
    'Get alex@identifind.com (or founder@identifind.com). Replace all Gmail references '
    'in pitch materials. Costs $6/user/month on Starter plan.',
    '$6/mo', 'workspace.google.com')
add_item_row(tbl1, False, 'HIGH', 'Point domain to your deployed app',
    'Add CNAME/A records in your domain registrar pointing to Vercel. '
    'Takes ~10 minutes — Vercel auto-provisions TLS.',
    'Free (Vercel)')
add_item_row(tbl1, False, 'HIGH', 'Register @identifind on all major social platforms',
    'Lock the handle on Twitter/X, LinkedIn (company page), Instagram, TikTok, and YouTube '
    'even if you won\'t use them yet — someone else will take them.',
    'Free')
add_item_row(tbl1, False, 'MEDIUM', 'Design a minimal logo / wordmark',
    'Use Figma (free) or hire a designer via Fiverr (~$50–150). '
    'You need this for the app, pitch deck, and email signature.',
    '$0–150')

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 2 — INFRASTRUCTURE UPGRADES
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '2.  Infrastructure Upgrades')

p2 = doc.add_paragraph()
p2.paragraph_format.space_before = Pt(6)
p2.paragraph_format.space_after  = Pt(8)
p2r = p2.add_run(
    'These prevent demo failures and show investors you operate like a professional engineering team.'
)
p2r.font.size = Pt(10); p2r.font.name = 'Arial'; p2r.font.color.rgb = MUTED

tbl2 = make_table(doc)
add_item_row(tbl2, False, 'CRITICAL', 'Upgrade Supabase to Pro tier',
    'Eliminates the free-tier project pausing (7-day inactivity = offline app). '
    'Adds daily backups, 8GB DB, 50GB bandwidth, and no cold-start pauses. '
    'Non-negotiable before any live investor demo.',
    '$25/mo', 'supabase.com/pricing')
add_item_row(tbl2, True,  'DONE',     'Prisma PgBouncer pooler configured',
    'DATABASE_URL uses port 6543 + pgbouncer=true. '
    'DIRECT_URL added to schema.prisma and .env for migrations.',
    'Free')
add_item_row(tbl2, False, 'CRITICAL', 'Fill in DIRECT_URL in .env',
    'Get from Supabase Dashboard → Settings → Database → Connection string → URI (port 5432). '
    'Required for prisma migrate deploy to succeed.',
    'Free', 'supabase.com/dashboard')
add_item_row(tbl2, False, 'HIGH', 'Deploy to Vercel (or confirm deployment)',
    'Connect GitHub repo to Vercel. Set all .env vars in Vercel → Settings → Environment Variables. '
    'Hobby tier is free; upgrade to Pro ($20/mo) before public launch for SLA and team features.',
    'Free → $20/mo', 'vercel.com')
add_item_row(tbl2, False, 'HIGH', 'Set NEXTAUTH_URL to production domain',
    'Change from http://localhost:3000 to https://identifind.com in your Vercel env vars. '
    'OAuth callbacks will silently fail without this.',
    'Free')
add_item_row(tbl2, False, 'MEDIUM', 'Enable Supabase Point-in-Time Recovery (PITR)',
    'Available on Pro tier. Gives you 7-day recovery window — important for a '
    'product handling personal identity data.',
    'Included in Pro')

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 3 — OBSERVABILITY & RELIABILITY
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '3.  Observability & Reliability')

p3 = doc.add_paragraph()
p3.paragraph_format.space_before = Pt(6)
p3.paragraph_format.space_after  = Pt(8)
p3r = p3.add_run(
    'Investors will ask "how do you know when something breaks?" '
    'These tools give you a credible answer and protect your users.'
)
p3r.font.size = Pt(10); p3r.font.name = 'Arial'; p3r.font.color.rgb = MUTED

tbl3 = make_table(doc)
add_item_row(tbl3, False, 'HIGH', 'Add Sentry for error tracking',
    'Install @sentry/nextjs. Captures unhandled exceptions with full stack traces, '
    'user context, and breadcrumbs. Free tier: 5K errors/month. '
    'Run: npm install @sentry/nextjs && npx @sentry/wizard@latest -i nextjs',
    'Free → $26/mo', 'sentry.io')
add_item_row(tbl3, False, 'HIGH', 'Set up uptime monitoring (BetterStack)',
    'Creates a public status page (status.identifind.com) and alerts you within 30s '
    'of downtime. Free tier covers 10 monitors. Looks extremely professional to investors.',
    'Free', 'betterstack.com/uptime')
add_item_row(tbl3, False, 'MEDIUM', 'Add PostHog for product analytics',
    'Self-hostable or cloud. Tracks user funnels, scan conversion, feature usage, '
    'and retention. Free up to 1M events/month. Investors will want this data at diligence.',
    'Free → $0.00031/event', 'posthog.com')
add_item_row(tbl3, False, 'MEDIUM', 'Configure Vercel Web Analytics',
    'One checkbox in Vercel dashboard. Shows page views, unique visitors, and countries '
    'at a glance — good for investor screenshots.',
    'Free on Hobby')
add_item_row(tbl3, False, 'LOW', 'Set up log aggregation (Axiom or Logtail)',
    'Capture server-side logs from Next.js API routes. '
    'Essential for debugging scan pipeline failures in production.',
    'Free tier available')

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 4 — TRANSACTIONAL EMAIL
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '4.  Transactional Email')

p4 = doc.add_paragraph()
p4.paragraph_format.space_before = Pt(6)
p4.paragraph_format.space_after  = Pt(8)
p4r = p4.add_run(
    'Users need email verification, password resets, and scan alert digests. '
    'Gmail\'s SMTP will get you rate-limited or spam-flagged at scale.'
)
p4r.font.size = Pt(10); p4r.font.name = 'Arial'; p4r.font.color.rgb = MUTED

tbl4 = make_table(doc)
add_item_row(tbl4, False, 'HIGH', 'Sign up for Resend (recommended)',
    'Built for Next.js — npm install resend, and you\'re sending in 5 minutes. '
    '3,000 emails/month free. Use noreply@identifind.com as the from address.',
    'Free → $20/mo', 'resend.com')
add_item_row(tbl4, False, 'HIGH', 'Configure SPF, DKIM, and DMARC DNS records',
    'Required for email deliverability. Resend walks you through adding these '
    'to your domain registrar. Without them, emails land in spam.',
    'Free (DNS config)')
add_item_row(tbl4, False, 'MEDIUM', 'Add Resend to NextAuth for magic-link / verification emails',
    'Replace the default nodemailer config with Resend\'s adapter. '
    'Docs: resend.com/docs/send-with-nextauth',
    'Free')

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 5 — API INTEGRATIONS (scan pipeline)
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '5.  API Integrations — Scan Pipeline')

p5 = doc.add_paragraph()
p5.paragraph_format.space_before = Pt(6)
p5.paragraph_format.space_after  = Pt(8)
p5r = p5.add_run(
    'Do these when you\'re ready to track expenses properly (LLC recommended first). '
    'Start with the free tiers — the pipeline already handles missing keys gracefully.'
)
p5r.font.size = Pt(10); p5r.font.name = 'Arial'; p5r.font.color.rgb = MUTED

tbl5 = make_table(doc)
add_item_row(tbl5, False, 'HIGH', 'HIBP (Have I Been Pwned) — breach data',
    'Most trusted breach database. ~$3.50/month. Get at haveibeenpwned.com/API/Key. '
    'Set HIBP_API_KEY in .env.',
    '$3.50/mo', 'haveibeenpwned.com')
add_item_row(tbl5, False, 'HIGH', 'People Data Labs — PII broker data',
    '100 free credits/month on free tier. Sign up at peopledatalabs.com. '
    'Set PDL_API_KEY in .env.',
    'Free → pay-per-use', 'peopledatalabs.com')
add_item_row(tbl5, False, 'HIGH', 'Refresh IntelX API key',
    'Current key returns 401 — regenerate at intelx.io/account?tab=developer. '
    'Set INTELX_API_KEY in .env.',
    'Free tier', 'intelx.io')
add_item_row(tbl5, False, 'MEDIUM', 'DeHashed — additional breach corpus',
    'Complements HIBP with additional leak data. API access at dehashed.com/api.',
    '~$5/mo', 'dehashed.com')
add_item_row(tbl5, False, 'MEDIUM', 'FullContact — PII enrichment',
    'Free tier available at platform.fullcontact.com. Set FULLCONTACT_API_KEY.',
    'Free tier', 'platform.fullcontact.com')
add_item_row(tbl5, False, 'LOW', 'OAuth provider keys (Google, GitHub, LinkedIn, Twitter, Facebook)',
    'Required for social account monitoring features. '
    'All blank in current .env — each needs its own developer app registered.',
    'Free')

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 6 — LEGAL & COMPLIANCE
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '6.  Legal & Compliance')

p6 = doc.add_paragraph()
p6.paragraph_format.space_before = Pt(6)
p6.paragraph_format.space_after  = Pt(8)
p6r = p6.add_run(
    'A product that handles personal identity data will be scrutinized during diligence. '
    'These items show investors you\'ve thought about risk.'
)
p6r.font.size = Pt(10); p6r.font.name = 'Arial'; p6r.font.color.rgb = MUTED

tbl6 = make_table(doc)
add_item_row(tbl6, False, 'CRITICAL', 'Form an LLC before signing up for paid services',
    'Register in Wyoming or Delaware (cheapest, most investor-friendly). '
    'Use Stripe Atlas ($500 one-time) or Clerky. All business expenses then flow through the entity.',
    '$50–500 one-time', 'stripe.com/atlas')
add_item_row(tbl6, False, 'CRITICAL', 'Open a business bank account',
    'Mercury (recommended for startups — no fees, no minimums, great API). '
    'Requires LLC to be formed first.',
    'Free', 'mercury.com')
add_item_row(tbl6, False, 'HIGH', 'Publish a Privacy Policy',
    'Required before collecting any user data. Use Termly or Iubenda to generate '
    'a CCPA/GDPR-compliant policy. Link from footer and auth screens.',
    '$10–30/mo', 'termly.io')
add_item_row(tbl6, False, 'HIGH', 'Publish Terms of Service',
    'Defines user obligations, limitations of liability, and acceptable use. '
    'Generate alongside Privacy Policy.',
    'Bundled with Privacy Policy')
add_item_row(tbl6, False, 'MEDIUM', 'Add cookie consent banner',
    'Required under GDPR for EU users. Termly or Cookiebot integrates with Next.js easily.',
    'Free tier available')
add_item_row(tbl6, False, 'LOW', 'Begin SOC 2 Type I preparation',
    'Not needed for pre-seed, but shows serious operational intent. '
    'Vanta ($800/mo) automates most of the evidence collection.',
    '$800/mo when ready', 'vanta.com')

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# SECTION 7 — INVESTOR-READINESS POLISH
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '7.  Investor-Readiness Polish')

p7 = doc.add_paragraph()
p7.paragraph_format.space_before = Pt(6)
p7.paragraph_format.space_after  = Pt(8)
p7r = p7.add_run(
    'The details that signal "serious founder" vs. "side project."'
)
p7r.font.size = Pt(10); p7r.font.name = 'Arial'; p7r.font.color.rgb = MUTED

tbl7 = make_table(doc)
add_item_row(tbl7, False, 'HIGH', 'Create a one-pager / teaser deck',
    'A 5–8 slide PDF you can email before a meeting: problem, solution, market, '
    'traction, ask. The Word pitch doc is your source of truth — convert it.',
    'Free')
add_item_row(tbl7, False, 'HIGH', 'Record a 3-minute product demo video',
    'Use Loom (free). Walk through a real scan with narration. '
    'Embed on the landing page and include the link in cold outreach.',
    'Free', 'loom.com')
add_item_row(tbl7, False, 'HIGH', 'Build a minimal landing page',
    'Explain what IdentiFind does, show the risk score UI, and capture emails. '
    'Use the existing Next.js app or a fast no-code tool (Framer, Webflow).',
    'Free → $20/mo')
add_item_row(tbl7, False, 'MEDIUM', 'Set up a waitlist / early-access sign-up',
    'Use Beehiiv, ConvertKit, or a simple Supabase-backed form. '
    '100+ waitlist signups is meaningful early traction for investors.',
    'Free tier available')
add_item_row(tbl7, False, 'MEDIUM', 'Get 5–10 beta users and collect testimonials',
    'Friends, family, or communities (r/privacy, r/netsec). '
    'Written quotes become social proof in the deck.',
    'Free')
add_item_row(tbl7, False, 'MEDIUM', 'Apply to YC, Techstars, or a local accelerator',
    'Applications are free. The process forces you to articulate the business clearly '
    'even if you don\'t get in. YC S2027 apps open ~Oct 2026.',
    'Free')
add_item_row(tbl7, False, 'LOW', 'Set up a Notion data room',
    'Shared folder with pitch deck, financials, cap table, and product docs. '
    'Professional investors expect a data room at first meeting.',
    'Free')

doc.add_paragraph()

# ════════════════════════════════════════════════════════════════════════════
# PRIORITY SUMMARY TABLE
# ════════════════════════════════════════════════════════════════════════════
add_section_heading(doc, '8.  90-Day Priority Sequence')

p8 = doc.add_paragraph()
p8.paragraph_format.space_before = Pt(6)
p8.paragraph_format.space_after  = Pt(8)
p8r = p8.add_run(
    'Recommended order if you\'re starting from zero today.'
)
p8r.font.size = Pt(10); p8r.font.name = 'Arial'; p8r.font.color.rgb = MUTED

seq_data = [
    ('Week 1',   'Form LLC, open Mercury bank account, register identifind.com, set up Google Workspace email.'),
    ('Week 2',   'Upgrade Supabase to Pro, fill in DIRECT_URL, deploy to Vercel with custom domain.'),
    ('Week 2',   'Set up Sentry, BetterStack status page, and Resend transactional email.'),
    ('Week 3',   'Publish Privacy Policy + Terms of Service, add cookie banner, register all social handles.'),
    ('Week 4',   'Activate HIBP and PDL API keys, refresh IntelX key, run first real end-to-end scan.'),
    ('Month 2',  'Build landing page, record demo video, launch waitlist. Start collecting beta users.'),
    ('Month 2',  'Begin OAuth provider setup (Google first), connect first real social account audit.'),
    ('Month 3',  'Mobile app Phase 2 (core screens). Apply to at least one accelerator program.'),
    ('Month 3',  'Set up PostHog analytics, produce first investor update with real metrics.'),
]

st = doc.add_table(rows=1, cols=2)
st.style = 'Table Grid'
for cell, txt in zip(st.rows[0].cells, ['Timeline', 'Action']):
    set_cell_bg(cell, '0F172A')
    r = cell.paragraphs[0].add_run(txt)
    r.bold = True; r.font.size = Pt(10); r.font.name = 'Arial'; r.font.color.rgb = WHITE
st.rows[0].cells[0].width = Inches(1.1)
st.rows[0].cells[1].width = Inches(6.4)

for timeline, action in seq_data:
    row = st.add_row()
    row.cells[0].width = Inches(1.1)
    row.cells[1].width = Inches(6.4)
    tr2 = row.cells[0].paragraphs[0].add_run(timeline)
    tr2.bold = True; tr2.font.size = Pt(10); tr2.font.name = 'Arial'; tr2.font.color.rgb = ACCENT
    ar = row.cells[1].paragraphs[0].add_run(action)
    ar.font.size = Pt(10); ar.font.name = 'Arial'

doc.add_paragraph()

add_divider(doc)

footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_p.paragraph_format.space_before = Pt(10)
fr = footer_p.add_run('IdentiFind  |  alexander.freshley@gmail.com  |  June 2026  |  Confidential')
fr.font.size = Pt(9); fr.font.name = 'Arial'; fr.font.color.rgb = MUTED

out = '/sessions/keen-nifty-tesla/mnt/IdentiFind/IdentiFind Launch Checklist.docx'
doc.save(out)
print(f'Saved: {out}')
